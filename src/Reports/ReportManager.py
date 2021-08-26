from datetime import timedelta

from docx import Document

from src.Movements.Models.Movement import Movement


def print_hi(name):
    # Use a breakpoint in the code line below to debug your script.
    print(f'Hi, {name}')  # Press Ctrl+F8 to toggle the breakpoint.


# Press the green button in the gutter to run the script.
class ReportManager:

    #Utente
    user = 'BENFATTO DANIELE'
    fiscal_code = '3yr24793dgq9ei08'
    contact = '0909504877'

    #document
    title = 'Il silmarillion'
    author = 'J.R.R. Tolkien'
    numInv = '874983749728'
    rack = 'X'
    shelf = 'A'
    position = 3
    coll = f"{rack} | {shelf} | {position}"
    dateIn = '20/20/2020'
    dateEnd = '50/50/2050'

    movement = Movement()

    def __init__(self, movement, newpath):

        self.newpath = newpath

        self.movement = movement

        #user
        self.user = self.movement.user.name + " " + self.movement.user.surname
        self.fiscal_code = self.movement.user.fiscal_code
        self.contact = self.movement.user.first_cellphone

        #document
        self.title = self.movement.item.title
        self.author = self.movement.item.author
        self.numInv = self.movement.item.id
        self.rack = self.movement.item.rack
        self.shelf = self.movement.item.shelf
        self.position = self.movement.item.position
        self.coll =  f"{self.rack} | {self.shelf} | {self.position}"
        self.dateIn = self.movement.timestamp
        self.dateEnd = self.movement.timestamp + timedelta(30)


    def set_report(self):
        doc = Document()

        #TITOLO REPORT
        doc.add_heading('MODULO DI PRESTITO', 0)
        core_properties = doc.core_properties




        par = doc.add_paragraph('\n\nBIBLIOTECA COMUNALE RECANATI\nCORSO PERSIANI 52 6100 RECANATI\n0719740021 biblioteca@comune.recanati.mc.it')
        #par.add_run('e un pò di grassetto').bold = True

        par = doc.add_paragraph(f"Il lettore  {self.user}\nCodice:     {self.fiscal_code}\nRecapiti:  {self.contact}")

        par = doc.add_paragraph('ottiene in prestito il documento con i seguenti dati di collocazione:')

        par = doc.add_paragraph(f"Titolo:    {self.title}\nAutore:  {self.author}\nN. inv:    {self.numInv}\nColloc:   {self.coll}\n")
        par.add_run(f"In prestito dal {self.dateIn} Scade il {self.dateEnd}").bold = True


        #doc.add_heading('Altro titolo', 1)
        #doc.add_paragraph('Quotazione', 'Intense Quote')

        #par = doc.add_paragraph('E infine una tabella')
        records = (
            ('ORARI DI APERTURA','',''),
            ('', 'Mattina', 'Pomeriggio'),
            ('Lun-Ven', '9:00 - 13:00', '15:30 - 18:30'),
            ('Sabato', '9:00 - 12:00', 'Chiuso')
        )


        table = doc.add_table(1, 3)
        #hdr_cells = table.rows[0].cells
        #[0].text = ''
        #hdr_cells[1].text = 'Mattina'
        #hdr_cells[2].text = 'Pomeriggio'

        for day, hours_M, hours_P in records:
            row_cells = table.add_row().cells
            row_cells[0].text = day
            row_cells[1].text = hours_M
            row_cells[2].text = hours_P

        par = doc.add_paragraph('\n\n\n\n\n                                                                             Firma\n\n\n\n'
                                '                                                    _______________________________________')


        doc.add_page_break()

        doc.save(f"{self.newpath}\Modulo di presito #{self.movement.id}.docx")

# See PyCharm help at https://www.jetbrains.com/help/pycharm/
